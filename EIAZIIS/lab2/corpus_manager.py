from unittest.mock import DEFAULT

import streamlit as st
import pandas as pd
import json
import re
from io import BytesIO
import fitz  # PyMuPDF вместо pypdf
from docx import Document
try:
    from striprtf.striprtf import rtf_to_text
except ImportError:
    rtf_to_text = None

import nltk
from razdel import tokenize
import pymorphy3
from performance_timer import PerformanceTimer

from constants import (
    MIN_YEAR, MAX_YEAR, DEFAULT_YEAR,
    POS_TRANSLATE, GRAMMEMES_TRANSLATE,
    TERMINOLOGY_TEXT, HELP_TEXT, MENU_ITEMS, CORPUS_COLUMNS
)

from sqlalchemy import create_engine, Column, Integer, String, Text
from sqlalchemy.orm import sessionmaker, declarative_base
from sqlalchemy import select, func, text

import os
import sys

# Force UTF-8 for all file system operations
os.environ["PYTHONUTF8"] = "1"

# Also set the locale to UTF-8
if sys.platform == "win32":
    os.environ["PYTHONIOENCODING"] = "utf-8"
    os.environ["LANG"] = "en_US.UTF-8"



nltk.download('punkt', quiet=True)

Base = declarative_base()

class CorpusDocument(Base):
    __tablename__ = 'documents'
    id = Column(Integer, primary_key=True)
    title = Column(String(255))
    author = Column(String(255))
    year = Column(Integer)
    type = Column(String(100))
    text = Column(Text)
    filename = Column(String(255))

class Token(Base):
    __tablename__ = 'tokens'
    id = Column(Integer, primary_key=True)
    doc_id = Column(Integer)
    wordform = Column(String(100))
    lemma = Column(String(100))
    pos_rus = Column(String(50))
    morph_rus = Column(Text)

class TagTranslator:
    """Класс для перевода морфологических тегов на русский язык."""

    POS_TRANSLATE = POS_TRANSLATE
    GRAMMEMES_TRANSLATE = GRAMMEMES_TRANSLATE

    def translate(self, tag_str):
        if not tag_str or tag_str == 'None':
            return ''
        parts = str(tag_str).replace(',', ', ').split()
        translated = []
        for part in parts:
            part = part.strip(', ')
            sub = [self.GRAMMEMES_TRANSLATE.get(p.strip(), "") for p in part.split(',')]
            translated.append(', '.join(sub))
        return ' '.join(translated).strip()

    def get_pos_rus(self, pos):
        return self.POS_TRANSLATE.get(pos, pos or "NONE")


class FileHandler:
    def extract_text(self, uploaded_file):
        name = uploaded_file.name.lower()
        data = uploaded_file.getvalue()
        if name.endswith('.txt'):
            return data.decode('utf-8', errors='ignore')
        elif name.endswith('.pdf'):
            doc = fitz.open(stream=data, filetype="pdf")
            text = '\n'.join(page.get_text() or '' for page in doc)
            doc.close()
            return text
        elif name.endswith('.docx'):
            doc = Document(BytesIO(data))
            return '\n'.join(p.text for p in doc.paragraphs)
        elif name.endswith('.rtf') and rtf_to_text:
            return rtf_to_text(data.decode('utf-8', errors='ignore'))
        return None


class TextParser:
    """Класс для разбора текста: токенизация и морфологический анализ."""

    def __init__(self):
        self.morph = pymorphy3.MorphAnalyzer(lang='ru')
        self.translator = TagTranslator()

    def analyze_text(self, text: str, doc_id: int):
        tokens = [token.text for token in tokenize(text) if re.match(r'^\w+$', token.text)]
        rows = []
        for token in tokens:
            if token.isdigit():  # числа исключены
                continue
            p = self.morph.parse(token)[0]
            pos = p.tag.POS or 'UNK'
            morph_tag = str(p.tag)
            rows.append({
                'doc_id': doc_id,
                'wordform': token.lower(),
                'lemma': p.normal_form.lower(),
                'pos_rus': self.translator.get_pos_rus(pos),
                'morph_rus': self.translator.translate(morph_tag),
            })
        return rows


import urllib.parse


class DataStorage:
    def __init__(self):
        import psycopg2
        from sqlalchemy import create_engine, text
        from sqlalchemy.orm import sessionmaker

        params = {
            'host': st.secrets['database']['host'],
            'port': st.secrets['database']['port'],
            'database': st.secrets['database']['database'],
            'user': st.secrets['database']['user'],
            'password': st.secrets['database']['password'],
            'client_encoding': 'utf8'
        }

        # Test raw psycopg2 connection first
        try:
            conn = psycopg2.connect(**params)
            conn.close()
            print("Raw psycopg2 connection OK")
        except Exception as e:
            print(f"Raw connection failed: {e}")
            raise

        # Now build SQLAlchemy engine using the same parameters
        # We'll create a custom connection function to bypass URL parsing
        def get_conn():
            return psycopg2.connect(**params)

        self.engine = create_engine('postgresql+psycopg2://', creator=get_conn)

        # Verify
        with self.engine.connect() as conn:
            conn.execute(text("SELECT 1"))

        Base.metadata.create_all(self.engine)
        self.Session = sessionmaker(bind=self.engine)

    def add_tokens(self, rows):
        if rows:
            with self.Session() as session:
                session.bulk_insert_mappings(Token, rows)
                session.commit()

    def add_document(self, text, metadata):
        doc = CorpusDocument(
            title=metadata['title'],
            author=metadata['author'],
            year=metadata['year'],
            type=metadata['type'],
            text=text,
            filename=metadata.get('filename')
        )
        with self.Session() as session:
            session.add(doc)
            session.commit()
            session.refresh(doc)
            return doc.id

    def update_document_text(self, doc_id, new_text):
        with self.Session() as session:
            doc = session.query(CorpusDocument).filter(CorpusDocument.id == doc_id).first()
            if doc:
                doc.text = new_text
                session.commit()

    def get_document_metadata(self, doc_id):
        with self.Session() as session:
            doc = session.query(CorpusDocument).filter(CorpusDocument.id == doc_id).first()
            if doc:
                return {
                    'title': doc.title,
                    'author': doc.author,
                    'year': doc.year,
                    'type': doc.type
                }
            return None

    def get_document_text(self, doc_id):
        with self.Session() as session:
            doc = session.query(CorpusDocument).filter(CorpusDocument.id == doc_id).first()
            return doc.text if doc else None

    def remove_tokens_for_doc(self, doc_id):
        with self.Session() as session:
            session.query(Token).filter(Token.doc_id == doc_id).delete()
            session.commit()

    def get_doc_list(self):
        query = select(CorpusDocument.id, CorpusDocument.title)
        with self.engine.connect() as conn:
            result = conn.execute(query).fetchall()
        return {row[0]: {'metadata': {'title': row[1]}} for row in result}

    # Статистические методы (SQL остается прежним, так как имена таблиц не менялись)
    def get_top_wordforms(self, limit=20):
        query = "SELECT wordform, COUNT(*) as frequency FROM tokens GROUP BY wordform ORDER BY frequency DESC LIMIT :limit"
        with self.engine.connect() as conn:
            result = conn.execute(text(query), {'limit': limit}).fetchall()
        return pd.DataFrame(result, columns=['Словоформа', 'Частота'])

    def get_top_lemmas(self, limit=20):
        query = "SELECT lemma, COUNT(*) as frequency FROM tokens GROUP BY lemma ORDER BY frequency DESC LIMIT :limit"
        with self.engine.connect() as conn:
            result = conn.execute(text(query), {'limit': limit}).fetchall()
        return pd.DataFrame(result, columns=['Лемма', 'Частота'])

    def get_pos_distribution(self):
        query = "SELECT pos_rus, COUNT(*) as count FROM tokens GROUP BY pos_rus ORDER BY count DESC"
        with self.engine.connect() as conn:
            result = conn.execute(text(query)).fetchall()
        return pd.DataFrame(result, columns=['Часть речи', 'Количество'])

    def get_top_morph(self, limit=15):
        query = "SELECT morph_rus, COUNT(*) as frequency FROM tokens WHERE morph_rus != '' AND morph_rus IS NOT NULL GROUP BY morph_rus ORDER BY frequency DESC LIMIT :limit"
        with self.engine.connect() as conn:
            result = conn.execute(text(query), {'limit': limit}).fetchall()
        return pd.DataFrame(result, columns=['Морфологические характеристики', 'Частота'])

    def get_doc_stats(self):
        query = """
        SELECT d.title as doc_title, COUNT(t.id) as count
        FROM documents d
        LEFT JOIN tokens t ON d.id = t.doc_id
        GROUP BY d.title
        ORDER BY count DESC
        """
        with self.engine.connect() as conn:
            result = conn.execute(text(query)).fetchall()
        return pd.DataFrame(result, columns=['Название документа', 'Количество словоформ'])

    def search_words(self, query):
        q = f"%{query.lower()}%"
        sql = """
        SELECT d.title as doc_title, t.wordform, t.lemma, t.pos_rus, t.morph_rus, d.author, d.year, d.type
        FROM tokens t
        JOIN documents d ON t.doc_id = d.id
        WHERE t.wordform ILIKE :q OR t.lemma ILIKE :q
        """
        with self.engine.connect() as conn:
            result = conn.execute(text(sql), {'q': q}).fetchall()
        return pd.DataFrame(result, columns=['Название документа', 'Словоформа', 'Лемма', 'Часть речи',
                                             'Морфологические характеристики', 'Автор', 'Год', 'Тип текста'])

    def search_phrases(self, phrase):
        q = f"%{phrase}%"
        sql = "SELECT id, title, text FROM documents WHERE text ILIKE :q"
        with self.engine.connect() as conn:
            return conn.execute(text(sql), {'q': q}).fetchall()

    def get_filtered_data(self, author=None, doc_type=None, year=None):
        where_clauses = []
        params = {}
        if author and author != "Все":
            where_clauses.append("d.author = :author")
            params['author'] = author
        if doc_type and doc_type != "Все":
            where_clauses.append("d.type = :type")
            params['type'] = doc_type
        if year and year != "Все":
            where_clauses.append("d.year = :year")
            params['year'] = year

        where = "WHERE " + " AND ".join(where_clauses) if where_clauses else ""
        sql = f"""
        SELECT d.title as doc_title, t.wordform, t.lemma, t.pos_rus, t.morph_rus, d.author, d.year, d.type
        FROM tokens t
        JOIN documents d ON t.doc_id = d.id
        {where}
        """
        with self.engine.connect() as conn:
            result = conn.execute(text(sql), params).fetchall()
        return pd.DataFrame(result, columns=['Название документа', 'Словоформа', 'Лемма', 'Часть речи',
                                             'Морфологические характеристики', 'Автор', 'Год', 'Тип текста'])

    def get_authors(self):
        query = select(func.distinct(CorpusDocument.author))
        with self.engine.connect() as conn:
            result = conn.execute(query).fetchall()
        return sorted([row[0] for row in result if row[0]])

    def get_types(self):
        query = select(func.distinct(CorpusDocument.type))
        with self.engine.connect() as conn:
            result = conn.execute(query).fetchall()
        return sorted([row[0] for row in result if row[0]])

    def get_years(self):
        query = select(func.distinct(CorpusDocument.year))
        with self.engine.connect() as conn:
            result = conn.execute(query).fetchall()
        return sorted([row[0] for row in result if row[0]], reverse=True)

    def save_to_json(self):
        docs_query = select(CorpusDocument)
        tokens_query = select(Token)
        with self.Session() as session:
            docs = session.execute(docs_query).scalars().all()
            tokens = session.execute(tokens_query).scalars().all()

        docs_dict = {doc.id: {'text': doc.text,
                              'metadata': {'title': doc.title, 'author': doc.author, 'year': doc.year, 'type': doc.type,
                                           'filename': doc.filename}} for doc in docs}
        corpus_records = [{
            'doc_id': t.doc_id, 'wordform': t.wordform, 'lemma': t.lemma, 'pos_rus': t.pos_rus, 'morph_rus': t.morph_rus
        } for t in tokens]

        data = {"docs": docs_dict, "corpus_df": corpus_records}
        return json.dumps(data, ensure_ascii=False, indent=2).encode('utf-8')

    def load_from_json(self, data):
        loaded = json.loads(data)
        old_to_new = {}
        with self.Session() as session:
            session.query(Token).delete()
            session.query(CorpusDocument).delete()
            session.commit()

            for old_id_str, doc_data in loaded['docs'].items():
                meta = doc_data['metadata']
                doc = CorpusDocument(title=meta['title'], author=meta['author'], year=meta['year'], type=meta['type'],
                                     text=doc_data['text'], filename=meta.get('filename'))
                session.add(doc)
                session.flush()
                old_to_new[int(old_id_str)] = doc.id

            token_rows = [{
                'doc_id': old_to_new[row['doc_id']], 'wordform': row['wordform'], 'lemma': row['lemma'],
                'pos_rus': row['pos_rus'], 'morph_rus': row['morph_rus']
            } for row in loaded['corpus_df'] if row['doc_id'] in old_to_new]

            if token_rows:
                session.bulk_insert_mappings(Token, token_rows)
            session.commit()

class View:

    def __init__(self, corpus_manager):
        self.corpus_manager = corpus_manager
        st.set_page_config(page_title="Корпусный менеджер", layout="wide")
        st.title("Корпусный менеджер текстов")

    def get_phrase_context(self, text: str, phrase: str, width: int = 90):
        pattern = re.compile(r'(.{0,' + str(width) + r'})(' + re.escape(phrase) + r')(.{0,' + str(width) + r'})', re.I)
        return pattern.findall(text)

    def display_menu(self):
        menu = st.sidebar.selectbox(
            "Меню",
            MENU_ITEMS
        )
        self.corpus_manager.handle_menu_selection(menu)

    def render_upload(self):
        st.header("Загрузка текстов в корпус")
        uploaded_files = st.file_uploader(
            "Выберите файлы (TXT, PDF, DOCX, RTF)",
            accept_multiple_files=True,
            type=['txt', 'pdf', 'docx', 'rtf']
        )

        st.subheader("Метаданные")
        col1, col2 = st.columns(2)
        with col1:
            author = st.text_input("Автор", "Неизвестен")
            title_base = st.text_input("Базовое название", "Документ")
        with col2:
            year = st.number_input("Год", MIN_YEAR, MAX_YEAR, DEFAULT_YEAR)
            doc_type = st.selectbox("Тип текста", ["Художественный", "Научный", "Публицистический", "Другой"])

        process_button = st.button(" Обработать и добавить в корпус", type="primary", key="process_files_button")
        return uploaded_files, author, title_base, year, doc_type, process_button

    def render_view_edit(self):
        st.header(" Просмотр и редактирование корпуса")
        docs = self.corpus_manager.storage.get_doc_list()
        if not docs:
            st.info("Корпус пока пустой")
            return None, None, False

        selected_id = st.selectbox("Выберите документ", list(docs.keys()),
                                   format_func=lambda x: f"{x} — {docs[x]['metadata']['title']}")
        text = self.corpus_manager.storage.get_document_text(selected_id)
        edited_text = st.text_area("Редактировать текст", text, height=400)
        save_button = st.button("Сохранить изменения", key="save_changes_button")
        return selected_id, edited_text, save_button

    def render_search_analysis(self):
        st.header("Поиск и анализ корпуса")
        # Check if corpus is empty
        with self.corpus_manager.storage.engine.connect() as conn:
            has_docs = conn.execute(text("SELECT COUNT(*) FROM documents")).scalar() > 0
        if not has_docs:
            st.warning("Сначала загрузите документы в раздел «Загрузка»")
            return

        tab_stats, tab_word, tab_phrase, tab_filter = st.tabs([
            "Общая статистика корпуса",
            "Поиск словоформ и лемм",
            "Поиск по фразе / кускам текста",
            "Фильтры по метаданным"
        ])

        with tab_stats:
            self.display_stats()

        with tab_word:
            query = self.render_word_search()
            self.display_word_search_results(query)

        with tab_phrase:
            phrase = self.render_phrase_search()
            self.display_phrase_search_results(phrase)

        with tab_filter:
            self.render_filters()

    def display_stats(self):
        st.subheader("Частотные характеристики по всему корпусу")

        col1, col2 = st.columns(2)
        with col1:
            st.write("**Топ-20 словоформ**")
            wf = self.corpus_manager.storage.get_top_wordforms()
            st.dataframe(wf, use_container_width=True)

            st.write("**Топ-20 лемм**")
            lem = self.corpus_manager.storage.get_top_lemmas()
            st.dataframe(lem, use_container_width=True)

        with col2:
            st.write("**Распределение по частям речи**")
            pos_df = self.corpus_manager.storage.get_pos_distribution()
            st.dataframe(pos_df, use_container_width=True)

        st.write("**Морфологические характеристики (топ-15)**")
        morph_df = self.corpus_manager.storage.get_top_morph()
        st.dataframe(morph_df, use_container_width=True)

        st.write("**Статистика по документам**")
        doc_stat = self.corpus_manager.storage.get_doc_stats()
        st.dataframe(doc_stat, use_container_width=True)

    def render_word_search(self):
        query = st.text_input("Введите слово или лемму для поиска")
        return query

    def display_word_search_results(self, query):
        if query:
            result = self.corpus_manager.storage.search_words(query)
            st.dataframe(result, use_container_width=True)

    def render_phrase_search(self):
        st.subheader("Поиск по произвольному тексту / словосочетанию")
        phrase = st.text_input("Введите фразу или кусок текста", placeholder="Пользовательский интерфейс")
        return phrase

    def display_phrase_search_results(self, phrase):
        if phrase:
            st.write(f"**Результаты поиска:** `{phrase}`")
            found_any = False
            results = self.corpus_manager.storage.search_phrases(phrase)
            for doc_id, title, text in results:
                contexts = self.get_phrase_context(text, phrase)
                if contexts:
                    found_any = True
                    st.subheader(f" {title}")
                    for left, match, right in contexts[:6]:
                        st.markdown(f"...{left}**{match}**{right}...")
                    st.divider()
            if not found_any:
                st.info("Фраза не найдена ни в одном документе корпуса.")

    def render_filters(self):
        st.subheader("Фильтрация по метаданным")
        authors = ["Все"] + self.corpus_manager.storage.get_authors()
        types_list = ["Все"] + self.corpus_manager.storage.get_types()
        years = ["Все"] + self.corpus_manager.storage.get_years()

        col1, col2, col3 = st.columns(3)
        with col1:
            sel_author = st.selectbox("Автор", authors)
        with col2:
            sel_type = st.selectbox("Тип текста", types_list)
        with col3:
            sel_year = st.selectbox("Год", years)

        filtered = self.corpus_manager.storage.get_filtered_data(sel_author, sel_type, sel_year)
        st.dataframe(filtered, use_container_width=True)

    def render_save_load(self):
        st.header("Сохранение и загрузка корпуса")
        col1, col2 = st.columns(2)
        save_button = False
        uploaded_json = None
        with col1:
            save_button = st.button("Сохранить корпус как JSON", key="save_json_button")
        with col2:
            uploaded_json = st.file_uploader("Загрузить сохранённый корпус", type='json')
        return save_button, uploaded_json

    def display_saved_json(self, data):
        st.download_button(
            label="Скачать corpus.json",
            data=data,
            file_name="corpus.json",
            mime="application/json",
            key="download_corpus"
        )

    def render_help(self):
        st.header("Система помощи")
        st.markdown(HELP_TEXT)

    def render_terminology(self):
        st.header("Терминологическая справка")
        st.markdown(TERMINOLOGY_TEXT)

    def render_sidebar_caption(self):
        st.sidebar.caption("Корпусный менеджер v0.999")


class CorpusManager:
    """Координирующий класс для управления всем процессом."""

    def __init__(self):
        self.storage = DataStorage()
        self.file_handler = FileHandler()
        self.parser = TextParser()
        self.view = View(self)

    def add_document(self, text: str, metadata: dict):
        doc_id = self.storage.add_document(text, metadata)
        rows = self.parser.analyze_text(text, doc_id)
        self.storage.add_tokens(rows)
        return doc_id

    def update_document(self, doc_id: int, new_text: str):
        metadata = self.storage.get_document_metadata(doc_id)
        if metadata:
            self.storage.remove_tokens_for_doc(doc_id)
            rows = self.parser.analyze_text(new_text, doc_id)
            self.storage.add_tokens(rows)
            self.storage.update_document_text(doc_id, new_text)

    def process_uploaded_files(self, uploaded_files, author, title_base, year, doc_type):
        PerformanceTimer.start()
        for file in uploaded_files:
            text = self.file_handler.extract_text(file)
            if text:
                metadata = {
                    "title": f"{title_base} — {file.name}",
                    "author": author,
                    "year": year,
                    "type": doc_type,
                    "filename": file.name
                }
                doc_id = self.add_document(text, metadata)
                st.success(f"{file.name} добавлен (ID {doc_id})")
        PerformanceTimer.stop()

    def handle_menu_selection(self, menu):
        if menu == "Загрузка и построение корпуса":
            self.handle_upload()
        elif menu == "Просмотр и редактирование корпуса":
            self.handle_view_edit()
        elif menu == "Поиск и анализ":
            self.handle_search_analysis()
        elif menu == "Сохранение / Загрузка":
            self.handle_save_load()
        elif menu == "Терминологическая справка":
            self.handle_terminology()
        else:
            self.handle_help()
        self.view.render_sidebar_caption()

    def handle_upload(self):
        uploaded_files, author, title_base, year, doc_type, process_button = self.view.render_upload()
        if uploaded_files and process_button:
            self.process_uploaded_files(uploaded_files, author, title_base, year, doc_type)

    def handle_view_edit(self):
        selected_id, edited_text, save_button = self.view.render_view_edit()
        if save_button and selected_id is not None:
            self.update_document(selected_id, edited_text)
            st.success("Изменения сохранены и корпус обновлён")

    def handle_search_analysis(self):
        self.view.render_search_analysis()

    def handle_save_load(self):
        save_button, uploaded_json = self.view.render_save_load()
        if save_button:
            data = self.storage.save_to_json()
            self.view.display_saved_json(data)
        if uploaded_json:
            data = uploaded_json.getvalue().decode('utf-8')
            self.storage.load_from_json(data)
            st.success("Корпус успешно загружен!")

    def handle_terminology(self):
        self.view.render_terminology()

    def handle_help(self):
        self.view.render_help()

    def run(self):
        self.view.display_menu()


if __name__ == "__main__":
    manager = CorpusManager()
    manager.run()